"""
Turn the Markdown policy sources into PDF and DOCX.

WHY THIS EXISTS
The corpus could have been indexed straight from Markdown. It is deliberately
not: generating real PDFs means Azure Document Intelligence has a genuine job in
the ingestion pipeline rather than a decorative one, and the demo can claim it
honestly. It also exercises the parsing quality we would face with a real
corpus — headings, tables and numbered sections — instead of a format that was
already machine-clean.

Markdown stays the source of truth. Nothing is ever hand-edited in PDF.

    python scripts/build_documents.py

Outputs land in data/generated/pdf and data/generated/docx.
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parent.parent
SOURCE_DIR = ROOT / "data" / "policies"
PDF_DIR = ROOT / "data" / "generated" / "pdf"
DOCX_DIR = ROOT / "data" / "generated" / "docx"

FOOTER = (
    "Sample content for demonstration purposes. Zuqah Technologies is a fictional company "
    "and this document describes no real organisation's policy."
)


# ---------------------------------------------------------------------------
# Parsing
# ---------------------------------------------------------------------------


@dataclass
class Block:
    """One renderable unit of a document."""

    kind: str  # h1 | h2 | h3 | para | bullet | table
    text: str = ""
    rows: list[list[str]] = field(default_factory=list)


@dataclass
class Policy:
    meta: dict[str, str]
    blocks: list[Block]

    @property
    def title(self) -> str:
        return self.meta.get("title", "Untitled")

    @property
    def slug(self) -> str:
        return self.meta.get("id", "untitled")


def parse_front_matter(text: str) -> tuple[dict[str, str], str]:
    """Split YAML-ish front matter from the body.

    Deliberately not a YAML parser: the front matter here is flat key/value only,
    and adding a dependency to read six keys would be the wrong trade.
    """
    if not text.startswith("---"):
        return {}, text

    _, raw, body = text.split("---", 2)
    meta: dict[str, str] = {}
    for line in raw.strip().splitlines():
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        meta[key.strip()] = value.strip().strip('"')
    return meta, body.strip()


def strip_inline(text: str) -> str:
    """Remove Markdown emphasis markers, keeping the words."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def parse_body(body: str) -> list[Block]:
    """Turn Markdown into an ordered list of blocks.

    Handles the subset the policy documents actually use: three heading levels,
    paragraphs, bullets, and pipe tables. Anything else is treated as a paragraph,
    which is the safe failure.
    """
    blocks: list[Block] = []
    lines = body.splitlines()
    index = 0
    paragraph: list[str] = []

    def flush_paragraph() -> None:
        if paragraph:
            blocks.append(Block("para", strip_inline(" ".join(paragraph).strip())))
            paragraph.clear()

    while index < len(lines):
        line = lines[index].rstrip()

        # Table: a pipe row followed by a separator row.
        if line.startswith("|") and index + 1 < len(lines) and set(lines[index + 1].replace("|", "").strip()) <= set("-: "):
            flush_paragraph()
            rows: list[list[str]] = []
            header = [strip_inline(c.strip()) for c in line.strip("|").split("|")]
            rows.append(header)
            index += 2  # skip the separator
            while index < len(lines) and lines[index].startswith("|"):
                rows.append([strip_inline(c.strip()) for c in lines[index].strip("|").split("|")])
                index += 1
            blocks.append(Block("table", rows=rows))
            continue

        if line.startswith("### "):
            flush_paragraph()
            blocks.append(Block("h3", strip_inline(line[4:])))
        elif line.startswith("## "):
            flush_paragraph()
            blocks.append(Block("h2", strip_inline(line[3:])))
        elif line.startswith("# "):
            flush_paragraph()
            blocks.append(Block("h1", strip_inline(line[2:])))
        elif line.startswith(("- ", "* ")):
            flush_paragraph()
            blocks.append(Block("bullet", strip_inline(line[2:])))
        elif re.match(r"^\d+\.\s", line):
            flush_paragraph()
            blocks.append(Block("bullet", strip_inline(line)))
        elif not line:
            flush_paragraph()
        elif line.startswith("---"):
            flush_paragraph()
        elif line.startswith("*Sample content"):
            flush_paragraph()  # the footer is added by the renderer
        else:
            paragraph.append(line)

        index += 1

    flush_paragraph()
    return blocks


def load_policies() -> list[Policy]:
    policies: list[Policy] = []
    for path in sorted(SOURCE_DIR.glob("*.md")):
        meta, body = parse_front_matter(path.read_text(encoding="utf-8"))
        policies.append(Policy(meta=meta, blocks=parse_body(body)))
    return policies


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def build_pdf(policy: Policy, out: Path) -> None:
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "PolicyTitle", parent=styles["Title"], fontSize=20, spaceAfter=6, alignment=TA_LEFT
    )
    meta_style = ParagraphStyle(
        "PolicyMeta", parent=styles["Normal"], fontSize=8.5, textColor=colors.grey, spaceAfter=18
    )
    h2_style = ParagraphStyle(
        "PolicyH2", parent=styles["Heading2"], fontSize=13, spaceBefore=16, spaceAfter=6
    )
    h3_style = ParagraphStyle(
        "PolicyH3", parent=styles["Heading3"], fontSize=11, spaceBefore=12, spaceAfter=4
    )
    body_style = ParagraphStyle(
        "PolicyBody", parent=styles["Normal"], fontSize=10, leading=14.5, spaceAfter=8
    )
    bullet_style = ParagraphStyle(
        "PolicyBullet", parent=body_style, leftIndent=16, bulletIndent=6, spaceAfter=4
    )
    footer_style = ParagraphStyle(
        "PolicyFooter", parent=styles["Normal"], fontSize=8, textColor=colors.grey, spaceBefore=24
    )

    story: list = []
    meta = policy.meta

    story.append(Paragraph(policy.title, title_style))
    story.append(
        Paragraph(
            f"Version {meta.get('version', '1.0')} &nbsp;|&nbsp; "
            f"Owner: {meta.get('owner', 'Unassigned')} &nbsp;|&nbsp; "
            f"Effective {meta.get('effective', '')} &nbsp;|&nbsp; "
            f"Next review {meta.get('review', '')}",
            meta_style,
        )
    )

    for block in policy.blocks:
        if block.kind == "h1":
            continue  # the title is already rendered from front matter
        if block.kind == "h2":
            story.append(Paragraph(block.text, h2_style))
        elif block.kind == "h3":
            story.append(Paragraph(block.text, h3_style))
        elif block.kind == "para":
            story.append(Paragraph(block.text, body_style))
        elif block.kind == "bullet":
            story.append(Paragraph(block.text, bullet_style, bulletText="•"))
        elif block.kind == "table":
            data = [[Paragraph(cell, body_style) for cell in row] for row in block.rows]
            table = Table(data, hAlign="LEFT", repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F1F5F9")),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(Spacer(1, 6))
            story.append(table)
            story.append(Spacer(1, 10))

    story.append(Paragraph(FOOTER, footer_style))

    SimpleDocTemplate(
        str(out),
        pagesize=LETTER,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=0.9 * inch,
        bottomMargin=0.9 * inch,
        title=policy.title,
        author="Zuqah Technologies",
        subject=meta.get("category", ""),
    ).build(story)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def build_docx(policy: Policy, out: Path) -> None:
    document = DocxDocument()
    meta = policy.meta

    heading = document.add_heading(policy.title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta_paragraph = document.add_paragraph()
    run = meta_paragraph.add_run(
        f"Version {meta.get('version', '1.0')}  |  "
        f"Owner: {meta.get('owner', 'Unassigned')}  |  "
        f"Effective {meta.get('effective', '')}  |  "
        f"Next review {meta.get('review', '')}"
    )
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    for block in policy.blocks:
        if block.kind == "h1":
            continue
        if block.kind == "h2":
            document.add_heading(block.text, level=1)
        elif block.kind == "h3":
            document.add_heading(block.text, level=2)
        elif block.kind == "para":
            document.add_paragraph(block.text)
        elif block.kind == "bullet":
            document.add_paragraph(block.text, style="List Bullet")
        elif block.kind == "table":
            table = document.add_table(rows=0, cols=len(block.rows[0]))
            table.style = "Light Grid Accent 1"
            for row in block.rows:
                cells = table.add_row().cells
                for cell, value in zip(cells, row):
                    cell.text = value

    footer = document.add_paragraph()
    footer_run = footer.add_run(FOOTER)
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    footer_run.italic = True

    document.save(str(out))


# ---------------------------------------------------------------------------


def main() -> int:
    policies = load_policies()

    if not policies:
        print(f"No Markdown sources found in {SOURCE_DIR}", file=sys.stderr)
        return 1

    PDF_DIR.mkdir(parents=True, exist_ok=True)
    DOCX_DIR.mkdir(parents=True, exist_ok=True)

    for policy in policies:
        pdf_path = PDF_DIR / f"{policy.slug}.pdf"
        docx_path = DOCX_DIR / f"{policy.slug}.docx"

        build_pdf(policy, pdf_path)
        build_docx(policy, docx_path)

        headings = sum(1 for b in policy.blocks if b.kind in ("h2", "h3"))
        tables = sum(1 for b in policy.blocks if b.kind == "table")
        words = sum(len(b.text.split()) for b in policy.blocks if b.text)

        print(
            f"{policy.slug:36} {words:5} words  "
            f"{headings:2} sections  {tables} tables  "
            f"pdf {pdf_path.stat().st_size // 1024:3} KB  "
            f"docx {docx_path.stat().st_size // 1024:3} KB"
        )

    print(f"\n{len(policies)} documents written to data/generated/")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
